import os
import tempfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

from app.reports.document import WordDocumentComposer


DOCUMENT_TITLE_PLACEHOLDER = "{{ document_title }}"
BODY_ANCHOR_PLACEHOLDER = "{{ body_anchor }}"
BODY_ANCHOR = "__TEST_CASE_REPORT_BODY__"


class ReportRenderError(RuntimeError):
    pass


def _validate_template_contract(template_path):
    document = Document(template_path)
    for placeholder, label in (
        (DOCUMENT_TITLE_PLACEHOLDER, "文档标题"),
        (BODY_ANCHOR_PLACEHOLDER, "正文锚点"),
    ):
        matches = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == placeholder
        ]
        if (
            len(matches) != 1
            or len(matches[0].runs) != 1
            or matches[0].runs[0].text != placeholder
        ):
            raise ReportRenderError(
                f"Word 模板必须用独立段落和单一文本块定义{label}占位符"
            )


def _enable_field_updates(document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _render_context(context, document_title):
    result = dict(context)
    metadata = dict(context["metadata"])
    metadata["document_name"] = document_title
    result["metadata"] = metadata
    return result


class WordReportRenderer:
    def __init__(self, profile):
        self._profile = profile

    def render(self, context, template, output_path):
        if not os.path.isfile(template.path):
            raise ReportRenderError("Word 模板不存在")

        _validate_template_contract(template.path)
        document_title = self._profile.document_title(context)
        context = _render_context(context, document_title)

        with tempfile.TemporaryDirectory(
            prefix="test-report-render-"
        ) as temp_dir:
            rendered_path = self._render_shell(
                template.path,
                document_title,
                temp_dir,
            )
            document = Document(rendered_path)
            anchor = self._find_anchor(document)
            composer = WordDocumentComposer(
                document,
                anchor,
                self._profile.theme,
            )

            for section in self._profile.create_sections():
                section.render(composer, context)

            anchor._element.getparent().remove(anchor._element)
            _enable_field_updates(document)
            document.core_properties.title = document_title
            document.save(output_path)

    @staticmethod
    def _render_shell(template_path, document_title, temp_dir):
        rendered_path = os.path.join(temp_dir, "rendered.docx")
        doc_template = DocxTemplate(template_path)
        doc_template.render(
            {
                "document_title": document_title,
                "body_anchor": BODY_ANCHOR,
            },
            autoescape=True,
        )
        doc_template.save(rendered_path)
        return rendered_path

    @staticmethod
    def _find_anchor(document):
        matches = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == BODY_ANCHOR
        ]
        if len(matches) != 1:
            raise ReportRenderError("Word 模板正文锚点渲染失败")
        return matches[0]
