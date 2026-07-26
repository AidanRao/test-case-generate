from dataclasses import dataclass

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass(frozen=True)
class HeadingStyle:
    size: float
    space_before: float
    space_after: float
    outline_level: int


@dataclass(frozen=True)
class CellMargins:
    top: int
    start: int
    bottom: int
    end: int


@dataclass(frozen=True)
class WordReportTheme:
    chinese_font: str
    latin_font: str
    body_size: float
    body_space_before: float
    body_space_after: float
    body_line_spacing: float
    metadata_size: float
    metadata_space_before: float
    metadata_space_after: float
    metadata_line_spacing: float
    step_label_size: float
    step_label_space_before: float
    step_label_space_after: float
    step_label_line_spacing: float
    table_font_size: float
    table_indent: int
    table_border_color: str
    table_header_fill: str
    table_cell_margins: CellMargins
    heading_styles: tuple[HeadingStyle, ...]

    def __post_init__(self):
        if len(self.heading_styles) != 4:
            raise ValueError("Word 报告主题必须定义四级标题样式")


def set_font(element, theme, size, bold=None):
    element.font.name = theme.latin_font
    element.font.size = Pt(size)
    if bold is not None:
        element.font.bold = bold
    run_properties = element._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), theme.latin_font)
    fonts.set(qn("w:hAnsi"), theme.latin_font)
    fonts.set(qn("w:cs"), theme.latin_font)
    fonts.set(qn("w:eastAsia"), theme.chinese_font)


def set_run_font(run, theme, size, bold=None):
    run.font.name = theme.latin_font
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), theme.latin_font)
    fonts.set(qn("w:hAnsi"), theme.latin_font)
    fonts.set(qn("w:cs"), theme.latin_font)
    fonts.set(qn("w:eastAsia"), theme.chinese_font)


def _clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _get_or_create_style(
    document,
    name,
    theme,
    size,
    bold,
    before,
    after,
    line_spacing=1.5,
):
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_font(style, theme, size=size, bold=bold)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    return style


def _configure_heading_style(style, theme, definition):
    set_font(style, theme, size=definition.size, bold=True)
    style.paragraph_format.space_before = Pt(definition.space_before)
    style.paragraph_format.space_after = Pt(definition.space_after)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True

    paragraph_properties = style._element.get_or_add_pPr()
    outline = paragraph_properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline)
    outline.set(qn("w:val"), str(definition.outline_level))


def configure_styles(document, theme):
    _get_or_create_style(
        document,
        "Report Body",
        theme,
        theme.body_size,
        False,
        theme.body_space_before,
        theme.body_space_after,
        line_spacing=theme.body_line_spacing,
    )
    _get_or_create_style(
        document,
        "Report Metadata",
        theme,
        theme.metadata_size,
        False,
        theme.metadata_space_before,
        theme.metadata_space_after,
        line_spacing=theme.metadata_line_spacing,
    )
    _get_or_create_style(
        document,
        "Report Step Label",
        theme,
        theme.step_label_size,
        True,
        theme.step_label_space_before,
        theme.step_label_space_after,
        line_spacing=theme.step_label_line_spacing,
    )
    for level, definition in enumerate(theme.heading_styles, start=1):
        _configure_heading_style(
            document.styles[f"Heading {level}"],
            theme,
            definition,
        )


def _set_cell_margins(cell, margins):
    cell_properties = cell._tc.get_or_add_tcPr()
    margins_element = cell_properties.first_child_found_in("w:tcMar")
    if margins_element is None:
        margins_element = OxmlElement("w:tcMar")
        cell_properties.append(margins_element)
    for name, value in (
        ("top", margins.top),
        ("start", margins.start),
        ("bottom", margins.bottom),
        ("end", margins.end),
    ):
        node = margins_element.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins_element.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width):
    cell_properties = cell._tc.get_or_add_tcPr()
    cell_width = cell_properties.first_child_found_in("w:tcW")
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        cell_properties.append(cell_width)
    cell_width.set(qn("w:w"), str(width))
    cell_width.set(qn("w:type"), "dxa")


class WordDocumentComposer:
    def __init__(self, document, anchor, theme):
        self.document = document
        self.anchor = anchor
        self.theme = theme
        configure_styles(document, theme)

    def add_heading(self, text, level):
        return self._add_paragraph(
            text,
            f"Heading {level}",
            keep_with_next=True,
        )

    def add_body(self, text):
        return self._add_paragraph(text, "Report Body")

    def add_label(self, text):
        return self._add_paragraph(
            text,
            "Report Step Label",
            keep_with_next=True,
        )

    def add_metadata(self, label, value):
        paragraph = self.anchor.insert_paragraph_before(
            style="Report Metadata"
        )
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(
            label_run,
            self.theme,
            size=self.theme.metadata_size,
            bold=True,
        )
        value_run = paragraph.add_run(str(value))
        set_run_font(
            value_run,
            self.theme,
            size=self.theme.metadata_size,
        )
        return paragraph

    def add_table(
        self,
        headers,
        rows,
        widths,
        centered_columns=(),
        merge_ranges=(),
    ):
        headers = tuple(headers)
        rows = tuple(rows)
        widths = tuple(widths)
        if len(widths) != len(headers):
            raise ValueError("表格列宽数量必须与表头数量一致")
        if any(len(row) != len(headers) for row in rows):
            raise ValueError("表格数据列数必须与表头数量一致")

        table = self.document.add_table(
            rows=1 + len(rows),
            cols=len(headers),
        )
        for column_index, header in enumerate(headers):
            self._format_cell(
                table.rows[0].cells[column_index],
                header,
                bold=True,
                centered=True,
                fill=self.theme.table_header_fill,
            )
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        header_properties.append(OxmlElement("w:tblHeader"))

        centered_columns = frozenset(centered_columns)
        for row_index, values in enumerate(rows, start=1):
            for column_index, value in enumerate(values):
                self._format_cell(
                    table.rows[row_index].cells[column_index],
                    value,
                    centered=column_index in centered_columns,
                )

        self._set_table_geometry(table, widths)
        for column, first_body_row, last_body_row in merge_ranges:
            if first_body_row < last_body_row:
                merged_cell = table.cell(first_body_row + 1, column).merge(
                    table.cell(last_body_row + 1, column)
                )
                merged_cell.text = ""
                self._format_cell(
                    merged_cell,
                    rows[first_body_row][column],
                    centered=column in centered_columns,
                )

        self.anchor._p.addprevious(table._tbl)
        spacer = self.anchor.insert_paragraph_before(style="Report Body")
        spacer.paragraph_format.space_after = Pt(2)
        return table

    def add_steps_table(self, steps):
        rows = [
            (index, step["step_desc"], step["expectation"])
            for index, step in enumerate(steps, start=1)
        ]
        return self.add_table(
            ("序号", "测试步骤", "预期结果"),
            rows,
            (720, 4140, 3450),
            centered_columns=(0,),
        )

    def _add_paragraph(
        self,
        text,
        style,
        keep_with_next=False,
    ):
        paragraph = self.anchor.insert_paragraph_before(style=style)
        run = paragraph.add_run(str(text))
        size = (
            paragraph.style.font.size.pt
            if paragraph.style.font.size
            else self.theme.body_size
        )
        set_run_font(run, self.theme, size=size)
        paragraph.paragraph_format.keep_with_next = keep_with_next
        return paragraph

    def _format_cell(
        self,
        cell,
        text,
        bold=False,
        centered=False,
        fill=None,
    ):
        paragraph = cell.paragraphs[0]
        _clear_paragraph(paragraph)
        paragraph.style = "Report Body"
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if centered
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        run = paragraph.add_run(str(text))
        set_run_font(
            run,
            self.theme,
            size=self.theme.table_font_size,
            bold=bold,
        )
        if fill:
            cell_properties = cell._tc.get_or_add_tcPr()
            shading = cell_properties.first_child_found_in("w:shd")
            if shading is None:
                shading = OxmlElement("w:shd")
                cell_properties.append(shading)
            shading.set(qn("w:fill"), fill)

    def _set_table_geometry(self, table, widths):
        table.autofit = False
        table_properties = table._tbl.tblPr
        table_width = table_properties.first_child_found_in("w:tblW")
        table_width.set(qn("w:w"), str(sum(widths)))
        table_width.set(qn("w:type"), "dxa")

        layout = table_properties.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table_properties.append(layout)
        layout.set(qn("w:type"), "fixed")

        indent = table_properties.first_child_found_in("w:tblInd")
        if indent is None:
            indent = OxmlElement("w:tblInd")
            table_properties.append(indent)
        indent.set(qn("w:w"), str(self.theme.table_indent))
        indent.set(qn("w:type"), "dxa")

        for grid_column, width in zip(table._tbl.tblGrid, widths):
            grid_column.set(qn("w:w"), str(width))

        borders = table_properties.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_properties.append(borders)
        for edge in (
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "4")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), self.theme.table_border_color)

        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            if row_properties.find(qn("w:cantSplit")) is None:
                row_properties.append(OxmlElement("w:cantSplit"))
            for cell, width in zip(row.cells, widths):
                _set_cell_width(cell, width)
                _set_cell_margins(
                    cell,
                    self.theme.table_cell_margins,
                )
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
