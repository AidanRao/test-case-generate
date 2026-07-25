import unittest

from docx import Document

from app.reports.document import WordDocumentComposer
from app.reports.sections import (
    DocumentOverviewSection,
    RequirementDetailsSection,
    RequirementOverviewSection,
)


CONTEXT = {
    "metadata": {
        "document_name": "示例系统测试用例文档",
        "project_name": "示例系统",
        "version": "V1.0",
        "compiled_date": "2026-07-25",
    },
    "project": {"id": "p-1", "code": "PRJ-01", "title": "示例系统"},
    "modules": [
        {
            "name": "认证",
            "requirements": [
                {
                    "id": "r-1",
                    "code": "REQ-001",
                    "title": "用户登录",
                    "type": "功能需求",
                    "content": "账号密码登录",
                    "testcases": [
                        {
                            "id": "tc-1",
                            "code": "TC-LOGIN-001",
                            "title": "登录成功",
                            "type": "功能测试",
                            "scenario_type": "正常流程用例",
                            "priority": "P0",
                            "test_target_desc": "验证登录",
                            "verify_method": "TESTING",
                            "steps": [
                                {
                                    "step_desc": "输入账号密码",
                                    "expectation": "登录成功",
                                }
                            ],
                        }
                    ],
                },
                {
                    "id": "r-2",
                    "code": "REQ-002",
                    "title": "用户注册",
                    "type": "功能需求",
                    "content": "注册账号",
                    "testcases": [],
                },
            ],
        }
    ],
    "summary": {
        "requirement_groups": [
            {
                "module": "认证",
                "requirements": [
                    {
                        "code": "REQ-001",
                        "title": "用户登录",
                        "type": "功能需求",
                        "testcase_count": 1,
                    },
                    {
                        "code": "REQ-002",
                        "title": "用户注册",
                        "type": "功能需求",
                        "testcase_count": 0,
                    },
                ],
            }
        ],
        "case_type_stats": [
            {"name": "功能测试", "count": 1, "percentage": "100.00%"}
        ],
        "priority_stats": [
            {"name": "P0", "count": 1, "percentage": "100.00%"}
        ],
        "coverage": [
            {
                "code": "REQ-001",
                "title": "用户登录",
                "testcase_codes": "TC-LOGIN-001",
                "testcase_count": 1,
                "status": "已覆盖",
            },
            {
                "code": "REQ-002",
                "title": "用户注册",
                "testcase_codes": "暂无",
                "testcase_count": 0,
                "status": "未覆盖",
            },
        ],
    },
}


class WordSectionTest(unittest.TestCase):
    def setUp(self):
        self.document = Document()
        self.anchor = self.document.add_paragraph("__ANCHOR__")
        self.composer = WordDocumentComposer(self.document, self.anchor)
        for section in (
            DocumentOverviewSection(),
            RequirementOverviewSection(),
            RequirementDetailsSection(),
        ):
            section.render(self.composer, CONTEXT)

    def table_for(self, headers):
        return next(
            table
            for table in self.document.tables
            if [cell.text for cell in table.rows[0].cells] == list(headers)
        )

    def test_renders_three_chapters_and_four_level_detail_hierarchy(self):
        paragraphs = {
            paragraph.text: paragraph.style.name
            for paragraph in self.document.paragraphs
        }
        self.assertEqual(paragraphs["一、文档概述"], "Heading 1")
        self.assertEqual(paragraphs["二、需求与用例概述"], "Heading 1")
        self.assertEqual(paragraphs["三、需求与测试用例明细"], "Heading 1")
        self.assertEqual(paragraphs["需求统计"], "Heading 2")
        self.assertEqual(paragraphs["用例统计"], "Heading 2")
        self.assertEqual(paragraphs["按用例类型统计"], "Heading 3")
        self.assertEqual(paragraphs["按优先级统计"], "Heading 3")
        self.assertEqual(paragraphs["需求覆盖分析"], "Heading 2")
        self.assertEqual(paragraphs["模块：认证"], "Heading 2")
        self.assertEqual(paragraphs["需求：用户登录"], "Heading 3")
        self.assertEqual(paragraphs["1. 登录成功"], "Heading 4")

    def test_renders_metadata_statistics_coverage_and_steps(self):
        metadata = self.table_for(("字段", "内容"))
        self.assertEqual(
            [[cell.text for cell in row.cells] for row in metadata.rows[1:]],
            [
                ["文档名称", "示例系统测试用例文档"],
                ["项目名称", "示例系统"],
                ["文档版本", "V1.0"],
                ["编制日期", "2026-07-25"],
            ],
        )

        requirement_stats = self.table_for(
            ("模块", "需求编号", "需求名称", "需求类型", "用例数量")
        )
        self.assertEqual(requirement_stats.rows[1].cells[0].text, "认证")
        self.assertIs(
            requirement_stats.rows[1].cells[0]._tc,
            requirement_stats.rows[2].cells[0]._tc,
        )

        case_types = self.table_for(("用例类型", "用例数量", "占比"))
        self.assertEqual(
            [cell.text for cell in case_types.rows[1].cells],
            ["功能测试", "1", "100.00%"],
        )
        priorities = self.table_for(("优先级", "用例数量", "占比"))
        self.assertEqual(
            [cell.text for cell in priorities.rows[1].cells],
            ["P0", "1", "100.00%"],
        )

        coverage = self.table_for(
            ("需求编号", "需求名称", "关联用例", "用例数量", "覆盖状态")
        )
        self.assertEqual(
            [cell.text for cell in coverage.rows[2].cells],
            ["REQ-002", "用户注册", "暂无", "0", "未覆盖"],
        )

        steps = self.table_for(("序号", "测试步骤", "预期结果"))
        self.assertEqual(
            [cell.text for cell in steps.rows[1].cells],
            ["1", "输入账号密码", "登录成功"],
        )
        self.assertTrue(
            any(
                paragraph.text == "测试用例：暂无"
                for paragraph in self.document.paragraphs
            )
        )

    def test_empty_report_keeps_headers_and_shows_no_data_messages(self):
        document = Document()
        anchor = document.add_paragraph("__ANCHOR__")
        composer = WordDocumentComposer(document, anchor)
        context = {
            "modules": [],
            "summary": {
                "requirement_groups": [],
                "case_type_stats": [],
                "priority_stats": [],
                "coverage": [],
            },
        }

        RequirementOverviewSection().render(composer, context)
        RequirementDetailsSection().render(composer, context)

        expected_empty_rows = {
            (
                "模块",
                "需求编号",
                "需求名称",
                "需求类型",
                "用例数量",
            ): ["暂无数据", "", "", "", ""],
            ("用例类型", "用例数量", "占比"): ["暂无数据", "", ""],
            ("优先级", "用例数量", "占比"): ["暂无数据", "", ""],
            (
                "需求编号",
                "需求名称",
                "关联用例",
                "用例数量",
                "覆盖状态",
            ): ["暂无数据", "", "", "", ""],
        }
        for headers, expected_row in expected_empty_rows.items():
            table = next(
                item
                for item in document.tables
                if [cell.text for cell in item.rows[0].cells]
                == list(headers)
            )
            self.assertEqual(
                [cell.text for cell in table.rows[1].cells],
                expected_row,
            )
        self.assertTrue(
            any(
                paragraph.text == "暂无需求"
                for paragraph in document.paragraphs
            )
        )


if __name__ == "__main__":
    unittest.main()
