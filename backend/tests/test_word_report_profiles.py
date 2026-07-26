from dataclasses import replace
from docx import Document
from docx.shared import Pt
import os
import tempfile
import unittest

from app.reports.profiles import (
    ReportProfileConfigurationError,
    STANDARD_WORD_REPORT_THEME,
    WordReportProfile,
    resolve_report_profile,
)
from app.reports.renderer import ReportRenderError, WordReportRenderer
from app.reports.sections import ReportSection
from app.reports.template_registry import ResolvedReportTemplate


CONTEXT = {
    "metadata": {
        "project_name": "示例系统",
        "version": "V1.0",
        "compiled_date": "2026-07-26",
    },
    "project": {
        "id": "project-1",
        "code": "PRJ-1",
        "title": "示例系统",
    },
    "modules": [],
    "summary": {
        "requirement_groups": [],
        "case_type_stats": [],
        "priority_stats": [],
        "coverage": [],
    },
}


def _template(path, title_count=1, body_count=1, title_size=24):
    document = Document()
    for _ in range(title_count):
        paragraph = document.add_paragraph()
        run = paragraph.add_run("{{ document_title }}")
        run.font.size = Pt(title_size)
        run.bold = True
    for _ in range(body_count):
        document.add_paragraph("{{ body_anchor }}")
    document.save(path)


def _resolved_template(path, profile_id="test"):
    return ResolvedReportTemplate(
        template_id="test-template",
        name="测试模板",
        path=path,
        profile_id=profile_id,
    )


class FirstSection(ReportSection):
    def render(self, composer, context):
        composer.add_heading("第一章节", 1)
        composer.add_body(context["metadata"]["document_name"])


class SecondSection(ReportSection):
    def render(self, composer, context):
        composer.add_heading("第二章节", 1)


class WordReportProfileTest(unittest.TestCase):
    def test_resolves_standard_profile_strictly(self):
        profile = resolve_report_profile("standard")

        self.assertEqual(profile.profile_id, "standard")
        self.assertEqual(
            profile.document_title(CONTEXT),
            "示例系统测试用例文档",
        )
        with self.assertRaises(ReportProfileConfigurationError):
            resolve_report_profile("missing")

    def test_profile_controls_section_order_title_and_theme(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = os.path.join(temp_dir, "template.docx")
            output_path = os.path.join(temp_dir, "output.docx")
            _template(template_path)
            heading = replace(
                STANDARD_WORD_REPORT_THEME.heading_styles[0],
                size=19,
            )
            theme = replace(
                STANDARD_WORD_REPORT_THEME,
                body_size=11,
                heading_styles=(
                    heading,
                    *STANDARD_WORD_REPORT_THEME.heading_styles[1:],
                ),
            )
            profile = WordReportProfile(
                profile_id="alternate",
                document_title_factory=lambda context: (
                    f"{context['project']['title']}精简报告"
                ),
                section_factories=(SecondSection, FirstSection),
                theme=theme,
            )

            WordReportRenderer(profile).render(
                CONTEXT,
                _resolved_template(template_path, profile.profile_id),
                output_path,
            )

            document = Document(output_path)
            texts = [paragraph.text for paragraph in document.paragraphs]
            self.assertLess(texts.index("第二章节"), texts.index("第一章节"))
            self.assertIn("示例系统精简报告", texts)
            first_heading = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text == "第一章节"
            )
            self.assertEqual(first_heading.runs[0].font.size.pt, 19)
            body_style = document.styles["Report Body"]
            self.assertEqual(body_style.font.size.pt, 11)

    def test_template_cover_formatting_is_preserved(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = os.path.join(temp_dir, "template.docx")
            output_path = os.path.join(temp_dir, "output.docx")
            _template(template_path, title_size=31)
            profile = WordReportProfile(
                profile_id="cover",
                document_title_factory=lambda context: "保留样式标题",
                section_factories=(),
                theme=STANDARD_WORD_REPORT_THEME,
            )

            WordReportRenderer(profile).render(
                CONTEXT,
                _resolved_template(template_path, profile.profile_id),
                output_path,
            )

            document = Document(output_path)
            title = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text == "保留样式标题"
            )
            self.assertEqual(title.runs[0].font.size.pt, 31)
            self.assertTrue(title.runs[0].bold)
            self.assertFalse(
                any(
                    paragraph.text == "__TEST_CASE_REPORT_BODY__"
                    for paragraph in document.paragraphs
                )
            )

    def test_template_contract_rejects_missing_or_duplicate_placeholders(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            profile = WordReportProfile(
                profile_id="contract",
                document_title_factory=lambda context: "测试标题",
                section_factories=(),
                theme=STANDARD_WORD_REPORT_THEME,
            )
            invalid_contracts = (
                ("missing-title.docx", 0, 1),
                ("duplicate-title.docx", 2, 1),
                ("missing-body.docx", 1, 0),
                ("duplicate-body.docx", 1, 2),
            )
            for filename, title_count, body_count in invalid_contracts:
                with self.subTest(filename=filename):
                    template_path = os.path.join(temp_dir, filename)
                    output_path = os.path.join(temp_dir, f"out-{filename}")
                    _template(
                        template_path,
                        title_count=title_count,
                        body_count=body_count,
                    )

                    with self.assertRaises(ReportRenderError):
                        WordReportRenderer(profile).render(
                            CONTEXT,
                            _resolved_template(
                                template_path,
                                profile.profile_id,
                            ),
                            output_path,
                        )

    def test_template_contract_rejects_placeholder_split_across_runs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = os.path.join(temp_dir, "split-placeholder.docx")
            output_path = os.path.join(temp_dir, "output.docx")
            document = Document()
            title = document.add_paragraph()
            title.add_run("{{ document_")
            title.add_run("title }}")
            document.add_paragraph("{{ body_anchor }}")
            document.save(template_path)
            profile = WordReportProfile(
                profile_id="contract",
                document_title_factory=lambda context: "测试标题",
                section_factories=(),
                theme=STANDARD_WORD_REPORT_THEME,
            )

            with self.assertRaises(ReportRenderError):
                WordReportRenderer(profile).render(
                    CONTEXT,
                    _resolved_template(template_path, profile.profile_id),
                    output_path,
                )


if __name__ == "__main__":
    unittest.main()
