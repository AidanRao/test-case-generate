from datetime import date
import io
import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook

from app.reports.template_registry import registered_template_ids, resolve_template


WORD_MIME = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def _table_for(document, headers):
    return next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells] == list(headers)
    )


def _table_rows(table):
    return [
        [cell.text for cell in row.cells]
        for row in table.rows
    ]


class WordReportExportTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmpdir.cleanup)
        self.environment = patch.dict(
            os.environ,
            {
                "DATA_DIR": self.tmpdir.name,
                "UNIPORTAL_SYNC_ENABLED": "false",
            },
            clear=False,
        )
        self.environment.start()
        self.addCleanup(self.environment.stop)

        from app import create_app

        self.app = create_app()
        report_date = patch("app.reports.context_builder.date")
        mocked_date = report_date.start()
        self.addCleanup(report_date.stop)
        mocked_date.today.return_value = date(2026, 7, 25)
        self.app.testing = True
        self.client = self.app.test_client()
        response = self.client.post(
            "/v1/projects",
            json={
                "code": "PRJ-WORD-01",
                "title": "Flight 2026飞控项目",
                "requirements": [
                    {
                        "module": "登录 Module 1",
                        "requirements": [
                            {
                                "title": "用户 Login 2026",
                                "type": "功能需求",
                                "code": "REQ-LOGIN-01",
                                "content": "用户输入 Account 123 后登录\n保留 $L_{stick}$ 原文。",
                            },
                            {
                                "title": "无用例需求",
                                "type": "",
                                "code": "REQ-EMPTY",
                                "content": "",
                            },
                        ],
                    },
                    {
                        "module": "审计",
                        "requirements": [
                            {
                                "title": "审计日志",
                                "type": "可靠性需求",
                                "code": "REQ-AUDIT",
                                "content": "记录审计事件。",
                            }
                        ],
                    },
                ],
            },
        )
        self.project_id = response.get_json()["data"]["id"]
        requirements = self.app.config["STORAGE"].list_requirements(self.project_id)
        login_requirement = requirements[0]
        audit_requirement = requirements[2]
        self.app.config["STORAGE"].add_testcases(
            self.project_id,
            login_requirement["id"],
            [
                {
                    "id": "TC-WORD-1",
                    "requirement_code": login_requirement["code"],
                    "requirement_id": login_requirement["id"],
                    "title": "正常 Login 01",
                    "code": "TC-LOGIN-01",
                    "type": "功能测试",
                    "scenario_type": "正常流程用例",
                    "priority": "P0",
                    "test_steps": [
                        {
                            "step_desc": "输入 Account 123",
                            "expectation": "登录 Success 200",
                        },
                        {
                            "step_desc": "查看首页",
                            "expectation": "显示欢迎信息",
                        },
                    ],
                    "test_target_desc": "验证登录流程",
                    "verify_method": "TESTING",
                }
            ],
        )
        self.app.config["STORAGE"].add_testcases(
            self.project_id,
            audit_requirement["id"],
            [
                {
                    "id": "TC-WORD-2",
                    "requirement_code": audit_requirement["code"],
                    "requirement_id": audit_requirement["id"],
                    "title": "审计空步骤",
                    "code": "TC-AUDIT-01",
                    "type": "功能测试",
                    "scenario_type": "异常场景用例",
                    "priority": "P1",
                    "test_steps": [],
                    "test_target_desc": "",
                    "verify_method": "",
                }
            ],
        )

    def test_excel_export(self):
        response = self.client.get(
            f"/v1/projects/{self.project_id}/testcases/export?format=xlsx"
        )
        self.addCleanup(response.close)

        self.assertEqual(response.status_code, 200)
        workbook = load_workbook(io.BytesIO(response.data))
        self.addCleanup(workbook.close)
        self.assertGreaterEqual(len(workbook.sheetnames), 1)

    def test_word_export_contains_report_hierarchy_and_step_table(self):
        response = self.client.get(
            f"/v1/projects/{self.project_id}/testcases/export"
            "?format=docx&template_id=default"
        )
        self.addCleanup(response.close)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.mimetype, WORD_MIME)
        self.assertIn(
            "PRJ-WORD-01-test-report.docx",
            response.headers["Content-Disposition"],
        )

        document = Document(io.BytesIO(response.data))
        texts = [paragraph.text for paragraph in document.paragraphs]
        expected_order = [
            "Flight 2026飞控项目测试用例文档",
            "一、文档概述",
            "二、需求与用例概述",
            "三、需求与测试用例明细",
            "模块：登录 Module 1",
            "需求：用户 Login 2026",
            "1. 正常 Login 01",
            "需求：无用例需求",
            "模块：审计",
            "需求：审计日志",
        ]
        positions = [texts.index(value) for value in expected_order]
        self.assertEqual(positions, sorted(positions))
        paragraphs_by_text = {
            paragraph.text: paragraph for paragraph in document.paragraphs
        }
        self.assertEqual(
            paragraphs_by_text["模块：登录 Module 1"].style.name,
            "Heading 2",
        )
        self.assertEqual(
            paragraphs_by_text["需求：用户 Login 2026"].style.name,
            "Heading 3",
        )
        self.assertEqual(
            paragraphs_by_text["1. 正常 Login 01"].style.name,
            "Heading 4",
        )
        self.assertEqual(
            [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style.name == "Heading 1"
            ],
            [
                "一、文档概述",
                "二、需求与用例概述",
                "三、需求与测试用例明细",
            ],
        )
        self.assertTrue(
            any("$L_{stick}$" in paragraph.text for paragraph in document.paragraphs)
        )
        self.assertTrue(
            any(paragraph.text == "测试用例：暂无" for paragraph in document.paragraphs)
        )
        self.assertTrue(
            any(paragraph.text == "测试步骤：暂无" for paragraph in document.paragraphs)
        )

        metadata_table = _table_for(document, ("字段", "内容"))
        self.assertEqual(
            _table_rows(metadata_table),
            [
                ["字段", "内容"],
                ["文档名称", "Flight 2026飞控项目测试用例文档"],
                ["项目名称", "Flight 2026飞控项目"],
                ["文档版本", "V1.0"],
                ["编制日期", "2026-07-25"],
            ],
        )

        requirement_table = _table_for(
            document,
            ("模块", "需求编号", "需求名称", "需求类型", "用例数量"),
        )
        self.assertEqual(
            [cell.text for cell in requirement_table.rows[1].cells],
            [
                "登录 Module 1",
                "REQ-LOGIN-01",
                "用户 Login 2026",
                "功能需求",
                "1",
            ],
        )
        self.assertIs(
            requirement_table.rows[1].cells[0]._tc,
            requirement_table.rows[2].cells[0]._tc,
        )

        case_type_table = _table_for(
            document,
            ("用例类型", "用例数量", "占比"),
        )
        self.assertEqual(
            _table_rows(case_type_table),
            [
                ["用例类型", "用例数量", "占比"],
                ["功能测试", "2", "100.00%"],
            ],
        )
        priority_table = _table_for(
            document,
            ("优先级", "用例数量", "占比"),
        )
        self.assertEqual(
            _table_rows(priority_table),
            [
                ["优先级", "用例数量", "占比"],
                ["P0", "1", "50.00%"],
                ["P1", "1", "50.00%"],
            ],
        )

        coverage_table = _table_for(
            document,
            ("需求编号", "需求名称", "关联用例", "用例数量", "覆盖状态"),
        )
        self.assertEqual(
            [cell.text for cell in coverage_table.rows[2].cells],
            ["REQ-EMPTY", "无用例需求", "暂无", "0", "未覆盖"],
        )
        self.assertEqual(
            [cell.text for cell in coverage_table.rows[1].cells],
            [
                "REQ-LOGIN-01",
                "用户 Login 2026",
                "TC-LOGIN-01",
                "1",
                "已覆盖",
            ],
        )

        step_table = _table_for(
            document,
            ("序号", "测试步骤", "预期结果"),
        )
        self.assertEqual(
            [cell.text for cell in step_table.rows[1].cells],
            ["1", "输入 Account 123", "登录 Success 200"],
        )
        self.assertEqual(
            [cell.text for cell in step_table.rows[2].cells],
            ["2", "查看首页", "显示欢迎信息"],
        )

    def test_word_structure_preserves_fonts_geometry_and_page_field(self):
        response = self.client.get(
            f"/v1/projects/{self.project_id}/testcases/export?format=docx"
        )
        self.addCleanup(response.close)
        document = Document(io.BytesIO(response.data))

        body_style = document.styles["Report Body"]
        r_fonts = body_style._element.rPr.rFonts
        self.assertEqual(r_fonts.get(qn("w:eastAsia")), "宋体")
        self.assertEqual(r_fonts.get(qn("w:ascii")), "Times New Roman")
        self.assertEqual(r_fonts.get(qn("w:hAnsi")), "Times New Roman")
        self.assertEqual(body_style.font.size.pt, 12)
        for style_name, outline_level in (
            ("Heading 1", "0"),
            ("Heading 2", "1"),
            ("Heading 3", "2"),
            ("Heading 4", "3"),
        ):
            heading_style = document.styles[style_name]
            outline = heading_style._element.pPr.find(qn("w:outlineLvl"))
            self.assertIsNotNone(outline)
            self.assertEqual(outline.get(qn("w:val")), outline_level)

        table = _table_for(
            document,
            ("序号", "测试步骤", "预期结果"),
        )
        widths = [
            int(grid_column.get(qn("w:w")))
            for grid_column in table._tbl.tblGrid
        ]
        self.assertEqual(widths, [720, 4140, 3450])
        self.assertEqual(
            table._tbl.tblPr.first_child_found_in("w:tblInd").get(qn("w:w")),
            "120",
        )
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        self.assertIsNotNone(header_properties.find(qn("w:tblHeader")))
        for row in table.rows:
            self.assertIsNotNone(
                row._tr.get_or_add_trPr().find(qn("w:cantSplit"))
            )

        footer_xml = "\n".join(
            section.footer._element.xml for section in document.sections
        )
        self.assertIn("PAGE", footer_xml)
        self.assertEqual(len(document.sections), 2)

    def test_invalid_export_parameters(self):
        missing_format = self.client.get(
            f"/v1/projects/{self.project_id}/testcases/export"
        )
        invalid_format = self.client.get(
            f"/v1/projects/{self.project_id}/testcases/export?format=pdf"
        )
        invalid_template = self.client.get(
            f"/v1/projects/{self.project_id}/testcases/export"
            "?format=docx&template_id=../../secret"
        )

        self.assertEqual(missing_format.status_code, 400)
        self.assertEqual(missing_format.get_json()["code"], 40001)
        self.assertEqual(invalid_format.status_code, 400)
        self.assertEqual(invalid_format.get_json()["code"], 40001)
        self.assertEqual(invalid_template.status_code, 400)
        self.assertEqual(invalid_template.get_json()["code"], 40001)

    def test_template_registry_is_allowlist(self):
        config = self.app.config["APP_CONFIG"]
        self.assertEqual(registered_template_ids(), ("default",))
        self.assertIsNotNone(resolve_template(config.base_dir, "default"))
        self.assertIsNone(resolve_template(config.base_dir, "../default"))


if __name__ == "__main__":
    unittest.main()
